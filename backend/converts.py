from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import subprocess

import datetime

from const_params import *


def add_zero_if_need(num: int) -> str:
    return str(num) if num > 9 else f"0{num}"


def convert_date(data_string: str) -> str:
    date = datetime.datetime.fromisoformat(data_string)
    return f"{add_zero_if_need(date.day)}.{add_zero_if_need(date.month)}.{date.year} {add_zero_if_need(date.hour)}:{add_zero_if_need(date.minute)}:{add_zero_if_need(date.second)}"


def save_as_json(data: dict):
    """ Формирование json-файла с результатом проверки """

    output_json_path = os.path.join(PATH_TO_DETECT_RESULTS, data['serial_number'],
                                    f"{data['serial_number']}.json")

    with open(file=output_json_path, mode="w", encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)


def get_json_file(serial_number: str) -> str:
    """ Получение пути до json-файла """

    return os.path.join(PATH_TO_DETECT_RESULTS, serial_number, f"{serial_number}.json")


def save_as_docx(serial_number: str):
    """ Оформление результатов диагностики в docx-файл """

    output_docx_path = os.path.join(PATH_TO_DETECT_RESULTS, serial_number,
                                    f"{serial_number}.docx")

    if (os.path.exists(output_docx_path)):
        return output_docx_path

    json_file_path = os.path.join(
        PATH_TO_DETECT_RESULTS, serial_number, f"{serial_number}.json")

    with open(json_file_path, 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)

    # Создать новый документ
    doc = Document()

    # Добавить заголовок
    doc.add_heading('Отчет об оценке состояния устройства', 0)

    serial_number = data['serial_number']

    # Добавить информацию об устройстве
    doc.add_heading('Информация об устройстве', level=1)
    doc.add_paragraph(f"Серийный номер: {data['serial_number']}")
    doc.add_paragraph(f"Дата проведения оценки: {convert_date(data['detect_defect_date'])}")
    doc.add_paragraph(f"Длительность проведения оценки: {data['detect_defect_working_time']} с")

    # Добавить итоговый результат
    doc.add_heading('Итоговый результат', level=1)
    doc.add_paragraph(f"Заключение о качестве: {data['detect_final_result']}")

    # Добавить сводную информацию о дефектах в виде таблицы
    doc.add_heading('Сводная информация об обнаруженных дефектах', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Добавить заголовки столбцов
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дефект'
    hdr_cells[1].text = 'Количество обнаружений'

    # Настроить стили для заголовков
    for cell in hdr_cells:
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

    # Добавить данные в таблицу
    for defect, count in data['all_defects'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = defect
        row_cells[1].text = str(count)

        # Центрировать текст во всех ячейках
        for cell in row_cells:
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавить информацию о дефектах по изображениям
    doc.add_heading('Дефекты, обнаруженные на изображениях', level=1)
    for img, defects in data['imgs_info'].items():
        doc.add_heading(f"Изображение: {img}", level=2)

        if defects:
            # Вставка изображения, если файл существует
            img_path = os.path.join(PATH_TO_DETECT_RESULTS,
                                    serial_number, OUTPUT_FOLDER_NAME, img)
            if os.path.exists(img_path):
                # Вставляем изображение шириной 2 дюйма
                doc.add_picture(img_path, width=Inches(2))
            for defect, count in defects.items():
                doc.add_paragraph(f"{defect}: {count} раз(а)")
        else:

            doc.add_paragraph("Дефекты не обнаружены")

    # Сохранить документ
    doc.save(output_docx_path)


def save_as_pdf(serial_number: str):
    """ Оформление результатов диагностики в pdf файл с использованием LibreOffice """

    # Пути к файлам
    pdf_output_path = os.path.join(PATH_TO_DETECT_RESULTS, serial_number, f"{serial_number}.pdf")
    docx_output_path = os.path.join(PATH_TO_DETECT_RESULTS, serial_number, f"{serial_number}.docx")

    # Проверяем, если PDF уже существует
    if os.path.exists(pdf_output_path):
        return pdf_output_path

    # Генерация DOCX, если его еще нет
    if not os.path.exists(docx_output_path):
        save_as_docx(serial_number)

    # Конвертация DOCX в PDF с использованием LibreOffice
    try:
        # Команда для конвертации
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(docx_output_path), docx_output_path],
            check=True
        )
    except subprocess.CalledProcessError as e:
        raise Exception(f"Ошибка при конвертации в PDF: {str(e)}")

    return pdf_output_path